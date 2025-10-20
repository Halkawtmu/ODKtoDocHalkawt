import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.set_page_config(page_title="XLSForm → Word (2‑Column)", page_icon="📄", layout="centered")

SECTION_COLOR_HEX = "1F4E79"  # dark blue

def set_cell_shading(cell, fill_hex=SECTION_COLOR_HEX):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def pick_label_value(row, df_columns, preferred_col=None):
    def val_ok(v):
        return v is not None and pd.notna(v) and str(v).strip() != ""
    # 1) Use preferred column if chosen
    if preferred_col and preferred_col in df_columns:
        v = row.get(preferred_col)
        if val_ok(v):
            return str(v).strip()
    # 2) Fallback to 'label'
    if 'label' in df_columns:
        v = row.get('label')
        if val_ok(v):
            return str(v).strip()
    # 3) Fallback to first 'label::<lang>'
    for col in df_columns:
        if str(col).lower().startswith('label::'):
            v = row.get(col)
            if val_ok(v):
                return str(v).strip()
    return ""

def extract_list_name(qtype, row):
    qt = (qtype or "").strip()
    parts = qt.split()
    if len(parts) >= 2 and (parts[0].startswith('select_one') or parts[0].startswith('select_multiple')):
        return parts[1]
    # fallback to a list_name column
    ln = row.get('list_name')
    if pd.notna(ln):
        return str(ln).strip()
    return None

def get_choices(choices_df, list_name, preferred_label_col=None):
    if not list_name:
        return []
    try:
        subset = choices_df[choices_df['list_name'].astype(str).str.strip() == str(list_name).strip()]
    except Exception:
        return []
    labels = []
    for _, r in subset.iterrows():
        lbl = pick_label_value(r, choices_df.columns, preferred_label_col)
        if lbl:
            labels.append(lbl)
    return labels

def build_doc(survey_df, choices_df, title_text="Community Policing - Satisfaction Survey", preferred_label_col=None):
    doc = Document()
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    section_font_size = Pt(12)
    question_font_size = Pt(11)

    # Build a 2-column table (left=question, right=choices/answer line)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    first_row_free = True

    def new_row():
        nonlocal first_row_free
        if first_row_free:
            first_row_free = False
            return table.rows[0]
        return table.add_row()

    def add_section_header(text):
        row = new_row()
        merged = row.cells[0].merge(row.cells[1])
        p = merged.paragraphs[0]
        p.text = ""
        run = p.add_run((text or "SECTION").upper())
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = section_font_size
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        set_cell_shading(merged, SECTION_COLOR_HEX)

    def add_question_row(q_text, choices=None):
        row = new_row()
        left = row.cells[0]
        right = row.cells[1]
        left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        p_left = left.paragraphs[0]
        p_left.text = ""
        r = p_left.add_run(q_text or "")
        r.bold = True
        r.font.size = question_font_size

        p_right = right.paragraphs[0]
        p_right.text = ""
        if choices and len(choices) > 0:
            for ch in choices:
                pr = right.add_paragraph(f"☐ {ch}")
                pr.paragraph_format.space_after = Pt(2)
        else:
            pr = right.paragraphs[0]
            pr.add_run("__________________________")
            pr.paragraph_format.space_after = Pt(6)

    # Loop survey rows
    for _, row in survey_df.iterrows():
        qtype = str(row.get('type', '')).strip()
        if not qtype:
            continue

        normalized = qtype.replace('_', ' ').lower()
        label = pick_label_value(row, survey_df.columns, preferred_label_col).strip()

        if normalized.startswith('begin group'):
            add_section_header(label or "SECTION")
            continue
        if normalized.startswith('end group'):
            continue
        if not label:
            continue

        list_name = extract_list_name(qtype, row)
        if list_name:
            choices = get_choices(choices_df, list_name, preferred_label_col)
            add_question_row(label, choices)
        else:
            add_question_row(label, None)

    # Set preferred column widths
    try:
        for cell in table.columns[0].cells:
            cell.width = Inches(4.5)  # Questions
        for cell in table.columns[1].cells:
            cell.width = Inches(2.5)  # Choices
    except Exception:
        pass

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

st.title("📄 XLSForm (KoBo/ODK) → Word Form (2‑Column)")
st.write("Upload your XLSForm (.xlsx) and download a Word form with two columns (Questions | Choices).")

uploaded = st.file_uploader("Upload XLSForm", type=["xlsx", "xls"])

if uploaded is not None:
    # Read Excel safely (case-insensitive sheet lookup)
    try:
        xls = pd.ExcelFile(uploaded)
        name_map = {n.lower(): n for n in xls.sheet_names}
        if 'survey' not in name_map or 'choices' not in name_map:
            st.error("Missing required sheets: 'survey' and/or 'choices'. Check your XLSForm.")
        else:
            survey_df = pd.read_excel(xls, sheet_name=name_map['survey'])
            choices_df = pd.read_excel(xls, sheet_name=name_map['choices'])

            # Detect available label columns from both sheets
            label_candidates = []
            for col in list(survey_df.columns) + list(choices_df.columns):
                c = str(col)
                if c.lower() == 'label' or c.lower().startswith('label::'):
                    label_candidates.append(c)
            label_candidates = sorted(set(label_candidates), key=lambda x: (x != 'label', x.lower()))

            with st.expander("Options", expanded=False):
                title_text = st.text_input("Document title", value="Community Policing - Satisfaction Survey")
                preferred_label = st.selectbox(
                    "Which label column should the app use?",
                    options=["Auto"] + label_candidates,
                    index=0,
                    help="If your form has multiple languages (label::<lang>), pick the one you want."
                )

            if st.button("Create Word (.docx)", type="primary"):
                with st.spinner("Generating Word document..."):
                    pref = None if preferred_label == "Auto" else preferred_label
                    doc_buf = build_doc(survey_df, choices_df, title_text=title_text, preferred_label_col=pref)
                out_name = f"{title_text.strip().replace(' ', '_')}_2col.docx"
                st.success("Your Word form is ready!")
                st.download_button(
                    "Download file",
                    data=doc_buf,
                    file_name=out_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    except Exception as e:
        st.error(f"Failed to read the Excel file. Details: {e}")
else:
    st.info("Drag & drop your XLSForm file here to get started.")